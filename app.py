from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
import requests
from dotenv import load_dotenv
import os
from datetime import datetime
import json
from functools import wraps
import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

app = Flask(__name__)
app.secret_key = os.urandom(24)

load_dotenv()

API_KEY = os.getenv("SYNCHROTEAM_API_KEY")
USERNAME = os.getenv("SYNCHRO_USERNAME")
PASSWORD = os.getenv("SYNCHRO_PASSWORD")

# Customer Locations (latitude, longitude)
CUSTOMER_LOCATIONS = {
    "Head Office": {"lat": -25.997590325139377, "lng": 28.062757040739807, "address": "Head Office"},
    "CVTS": {"lat": -25.97557331702371, "lng": 28.076559879538017, "address": "CVTS"},
    "ACME Corp": {"lat": -26.194555083288634, "lng": 28.03197028938919, "address": "ACME Corp"},
}

# Cache storage
cache = {
    'data': None,
    'last_sync': None,
    'last_sync_str': 'Never'
}

def haversine_distance(lat1, lon1, lat2, lon2):
    """Calculate the great circle distance between two points on the earth"""
    lat1, lon1, lat2, lon2 = map(math.radians, [lat1, lon1, lat2, lon2])
    dlon = lon2 - lon1
    dlat = lat2 - lat1
    a = math.sin(dlat/2)**2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon/2)**2
    c = 2 * math.asin(math.sqrt(a))
    r = 6371
    return c * r

def calculate_distances(jobs):
    """Calculate distance from Head Office for each job"""
    head_office = CUSTOMER_LOCATIONS["Head Office"]
    
    for job in jobs:
        customer_name = job.get('customer', '')
        if customer_name in CUSTOMER_LOCATIONS:
            customer_loc = CUSTOMER_LOCATIONS[customer_name]
            distance = haversine_distance(
                head_office["lat"], head_office["lng"],
                customer_loc["lat"], customer_loc["lng"]
            )
            job['distance'] = round(distance, 2)
            job['customer_lat'] = customer_loc["lat"]
            job['customer_lng'] = customer_loc["lng"]
        else:
            job['distance'] = None
            job['customer_lat'] = None
            job['customer_lng'] = None
    
    return jobs

def require_login(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session or not session['logged_in']:
            flash('Please login to access the dashboard.', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

def set_cell_background(cell, color_hex):
    """Set background color for a table cell"""
    try:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'solid')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    except:
        pass  

def fetch_jobs_data():
    """Fetch jobs data from Synchroteam API"""
    try:
        url = "https://ws.synchroteam.com/api/v3/job/list?"
        querystring = {"dateFrom": "2026-01-01", "dateTo": "2026-07-07", "pageSize": 100}
        
        headers = {
            'authorization': API_KEY,
            'accept': "text/json",
            'content-type': "application/json",
            'cache-control': "no-cache"
        }
        
        response = requests.get(url, headers=headers, params=querystring)
        response.raise_for_status()
        list_data = response.json()
        
        jobs = []
        
        for record in list_data.get("data", []):
            record_id = record.get("id")
            
            if record_id:
                detail_url = "https://ws.synchroteam.com/api/v3/job/details?"
                detail_params = {"id": record_id}
                
                detail_response = requests.get(detail_url, headers=headers, params=detail_params)
                detail_response.raise_for_status()
                detail_data = detail_response.json()
                
                job_info = {
                    'num': detail_data.get("num", "N/A"),
                    'priority': detail_data.get("priority", "N/A"),
                    'jobStatus': detail_data.get("status", "N/A"),
                    'description': detail_data.get("description", ""),
                    'jobType': detail_data.get("type", {}).get("name", "N/A"),
                    'customer': detail_data.get("customer", {}).get("name", "N/A"),
                    'technician': detail_data.get("technician", {}).get("name", "Unassigned"),
                }
                
                jobs.append(job_info)
        
        jobs = calculate_distances(jobs)
        
        return {
            'jobs': jobs,
            'total': len(jobs),
            'success': True
        }
        
    except requests.exceptions.RequestException as e:
        return {
            'success': False,
            'error': f'API Error: {str(e)}'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Unexpected error: {str(e)}'
        }

@app.route('/')
def index():
    if 'logged_in' in session and session['logged_in']:
        return redirect(url_for('dashboard'))
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username')
    password = request.form.get('password')
    remember = request.form.get('remember')
    
    if username == USERNAME and password == PASSWORD:
        session['logged_in'] = True
        session['username'] = username
        
        if remember:
            session.permanent = True
        
        flash('Login successful!', 'success')
        return redirect(url_for('dashboard'))
    else:
        flash('Invalid username or password. Please try again.', 'error')
        return redirect(url_for('index'))

@app.route('/dashboard')
@require_login
def dashboard():
    global cache
    
    if cache['data'] is None:
        result = fetch_jobs_data()
        if result['success']:
            cache['data'] = result
            cache['last_sync'] = datetime.now()
            cache['last_sync_str'] = cache['last_sync'].strftime('%Y-%m-%d %H:%M:%S')
            flash(f'Successfully loaded {len(result["jobs"])} jobs!', 'success')
        else:
            flash(result['error'], 'error')
            cache['data'] = {'jobs': [], 'total': 0}
    else:
        flash(f'Showing cached data from {cache["last_sync_str"]}', 'info')
    
    return render_template('dashboard.html', 
                         data=cache['data'],
                         last_sync=cache['last_sync_str'],
                         customer_locations=CUSTOMER_LOCATIONS,
                         head_office=CUSTOMER_LOCATIONS["Head Office"])

@app.route('/download-report')
@require_login
def download_report():
    """Generate and download a Word document report"""
    global cache
    
    if cache['data'] is None or not cache['data'].get('jobs'):
        flash('No data available to generate report.', 'error')
        return redirect(url_for('dashboard'))
    
    jobs = cache['data']['jobs']
    
    # Create a new Document
    doc = Document()
    
    # --- HEADER ---
    # Add a title with background using a table
    header_table = doc.add_table(rows=1, cols=1)
    header_table.autofit = False
    header_table.rows[0].cells[0].text = '📊 Executive Dashboard Report'
    header_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Style the header text
    run = header_table.rows[0].cells[0].paragraphs[0].runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Set background color for header
    set_cell_background(header_table.rows[0].cells[0], '667eea')
    
    doc.add_paragraph('')  
    
    # --- SUBTITLE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Total Jobs: {len(jobs)}')
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph('')
    
    # --- SUMMARY STATISTICS ---
    doc.add_heading('Summary Statistics', level=2)
    
    # Calculate stats
    completed_count = sum(1 for j in jobs if 'complete' in j.get('jobStatus', '').lower() or 'done' in j.get('jobStatus', '').lower())
    high_priority_count = sum(1 for j in jobs if j.get('priority', '').lower() in ['high', 'urgent'])
    unique_customers = set(j.get('customer', '') for j in jobs)
    
    # Create stats table
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.style = 'Table Grid'
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = stats_table.rows[0].cells
    headers = ['Total Jobs', 'Completed', 'High Priority', 'Unique Customers']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set background color for headers
        set_cell_background(hdr_cells[i], 'e8e8e8')
    
    # Data row
    row_cells = stats_table.rows[1].cells
    row_cells[0].text = str(len(jobs))
    row_cells[1].text = str(completed_count)
    row_cells[2].text = str(high_priority_count)
    row_cells[3].text = str(len(unique_customers))
    
    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph('')
    
    # --- JOBS TABLE ---
    doc.add_heading('Job Details', level=2)
    
    # Create table with headers
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr_cells = table.rows[0].cells
    headers = ['Job #', 'Customer', 'Status', 'Type', 'Description', 'Priority', 'Technician']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set background color for table headers
        set_cell_background(hdr_cells[i], 'e8e8e8')
    
    # Add data rows (limit to 50 for readability)
    display_jobs = jobs[:50]
    for job in display_jobs:
        row_cells = table.add_row().cells
        row_cells[0].text = str(job.get('num', 'N/A'))
        row_cells[1].text = str(job.get('customer', 'N/A'))
        row_cells[2].text = str(job.get('jobStatus', 'N/A'))
        row_cells[3].text = str(job.get('jobType', 'N/A'))
        
        desc = str(job.get('description', ''))
        row_cells[4].text = desc[:50] + ('...' if len(desc) > 50 else '')
        
        # Highlight priority
        priority = str(job.get('priority', 'N/A'))
        row_cells[5].text = priority
        
        row_cells[6].text = str(job.get('technician', 'Unassigned'))
        
        # Center align all cells
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add note if more than 50 jobs
    if len(jobs) > 50:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'* Showing first 50 of {len(jobs)} jobs')
        run.font.italic = True
        run.font.size = Pt(10)
    
    doc.add_paragraph('')
    
    # --- CUSTOMER LOCATIONS ---
    doc.add_heading('Customer Locations', level=2)
    
    head_office = CUSTOMER_LOCATIONS["Head Office"]
    
    # Create a table for customer locations
    loc_table = doc.add_table(rows=1, cols=3)
    loc_table.style = 'Table Grid'
    loc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    loc_hdr = loc_table.rows[0].cells
    loc_headers = ['Customer', 'Distance from Head Office (km)', 'Job Count']
    for i, header in enumerate(loc_headers):
        loc_hdr[i].text = header
        loc_hdr[i].paragraphs[0].runs[0].font.bold = True
        loc_hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(loc_hdr[i], 'e8e8e8')
    
    # Add customer data
    for name, loc in CUSTOMER_LOCATIONS.items():
        if name != 'Head Office':
            distance = haversine_distance(
                head_office["lat"], head_office["lng"],
                loc["lat"], loc["lng"]
            )
            job_count = sum(1 for j in jobs if j.get('customer') == name)
            
            row_cells = loc_table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = f'{distance:.2f}'
            row_cells[2].text = str(job_count)
            
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # --- FOOTER ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('-' * 60)
    run.font.size = Pt(10)
    
    doc.add_paragraph('')
    p = doc.add_paragraph('This report was generated automatically by the Executive Dashboard System.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    p = doc.add_paragraph('© 2026 Executive Dashboard • Powered by Synchroteam')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Generate filename
    filename = f'Executive_Dashboard_Report_{datetime.now().strftime("%Y-%m-%d")}.docx'
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/refresh')
@require_login
def refresh_data():
    """Endpoint to manually refresh the data"""
    global cache
    
    result = fetch_jobs_data()
    
    if result['success']:
        cache['data'] = result
        cache['last_sync'] = datetime.now()
        cache['last_sync_str'] = cache['last_sync'].strftime('%Y-%m-%d %H:%M:%S')
        flash(f'Data refreshed! Loaded {len(result["jobs"])} jobs.', 'success')
    else:
        flash(result['error'], 'error')
    
    return redirect(url_for('dashboard'))

@app.route('/api/refresh')
@require_login
def api_refresh():
    """API endpoint for AJAX refresh"""
    global cache
    
    result = fetch_jobs_data()
    
    if result['success']:
        cache['data'] = result
        cache['last_sync'] = datetime.now()
        cache['last_sync_str'] = cache['last_sync'].strftime('%Y-%m-%d %H:%M:%S')
        return jsonify({
            'success': True,
            'message': f'Refreshed {len(result["jobs"])} jobs',
            'last_sync': cache['last_sync_str'],
            'total': len(result['jobs'])
        })
    else:
        return jsonify({
            'success': False,
            'error': result['error']
        }), 500

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'success')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)